# pip install python-docx 
from docx import Document
import os

# import pip
# package_names=['docx'] #packages to install
# pip.main(['install'] + package_names)
#pip.main(['install'] + package_names + ['--upgrade']) 

os.getcwd()
os.chdir('C:\\Users\\A485205\\Desktop\\Hackathon\\openai-hackathon')

# .venv\Scripts\activate

# Extract number and name headings       
def extract_number_heading(doc_path: str):
    '''
    Input: path of a docx file 
    Output: a list of heading names and their number headings corresponding chapters, sections, subsections,... 
    '''
    doc = Document(doc_path)
    # Extract chapters, section, subsection numbers
    heading_numbers=[]
    heading_name=[]

    for paragraph in doc.paragraphs:
        # Check if the paragraph is a heading
        if paragraph.style.name.startswith('Heading'):            # Heading for all chapters, sections, subsections,.., Heading 1 for chapter, Heading 2 for subsections, Heading 3 for subsubsections
            # Save heading_numbers       
            heading_numbers.append(paragraph.style.name.split()[1])
            heading_name.append((paragraph.text))
    #print(heading_numbers)
    
    # Map heading_numbers to heading sections
    heading_sections_result = []
    chapter = 1
    for i in range(len(heading_numbers)):
        if heading_numbers[i] =='1':
            heading_sections_result.append(str(chapter))
            chapter+=1
        else:
            if int(heading_numbers[i]) > int(heading_numbers[i-1]):

                if int(heading_numbers[i]) == 2:
                    heading_sections_result.append(heading_sections_result[-1]+'.1')
                else:
                    heading_sections_result.append(heading_sections_result[-1] + '.1'*(int(heading_numbers[i])-2))


            elif int(heading_numbers[i]) == int(heading_numbers[i-1]):
                u = heading_sections_result[-1].split('.')
                u[-1] = str(int(u[-1])+1)
                heading_sections_result.append('.'.join(u))
       
            elif int(heading_numbers[i]) < int(heading_numbers[i-1]):
                u = heading_sections_result[-1].split('.')
                u = u[:int(heading_numbers[i])]
                u[-1] = str(int(u[-1])+1)
                heading_sections_result.append('.'.join(u))
    
    # Combine heading_name and heading_sections_result 
    total_heading = []
    for i in range(len(heading_numbers)):
        total_heading.append(heading_sections_result[i] + ' ' + heading_name[i])

    return total_heading

# Splitting documents
def save_content_under_headings(doc_path, folder_name):

    errors = []
    # Load the Word document
    doc = Document(doc_path)
    
    # Create a new folder
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    
    current_heading = None
    current_doc = None
    
    doc_name = doc_path.split('/')[-1]
    doc_name = '['+ ''.join(doc_name.split(".")[0]) +']' + ' '
    heading_name_list = extract_number_heading(doc_path)
    k=-1

    for paragraph in doc.paragraphs:
        # Check if the paragraph is a heading
        if paragraph.style.name.startswith('Heading'):            # Heading for all chapters, sections, subsections,.., Heading 1 for chapter, Heading 2 for subsections, Heading 3 for subsubsections
            # Save the previous section if there is one           
            #print(paragraph.style.name.split()[1])

            current_heading = heading_name_list[k]
            current_heading = current_heading.replace('/', '-').replace('\\', '-').replace('\t','')
            k+=1

            if current_heading and current_doc:
                file_path = os.path.join(folder_name, f"{doc_name}{current_heading}.docx")
                #file_path = f"{folder_name}\\{doc_name}{current_heading}.docx"
                print(file_path)
                try:
                    current_doc.save(file_path)
                except:
                    print(file_path)
                    errors.append(current_heading)
                    pass
            
            # Reset for the next section
            current_heading = heading_name_list[k]
            current_heading = current_heading.replace('/', '-').replace('\\', '-').replace('\t','')
            current_doc = Document()
            current_doc.add_heading(current_heading, level=0)
        else:
            # Accumulate content under the current heading
            if current_doc is not None:
                current_doc.add_paragraph(paragraph.text)
    
    # Don't forget to save the last section
    if current_heading and current_doc:
        current_heading = heading_name_list[k]
        current_heading = current_heading.replace('/', '-').replace('\\', '-').replace('\t','')
        file_path = os.path.join(folder_name, f"{doc_name}{current_heading}.docx")
        #file_path = f"{folder_name}\\{doc_name}{current_heading}.docx"
        
        try:
            current_doc.save(file_path)
        except:
            errors.append(current_heading)
            print(file_path)
            pass
    
    if len(errors)>1:
        file = open(os.path.join(folder_name, sections_errors.txt), 'w')
        file.write('\n\n'.join(errors))
        file.close() 
    
    # End of the function
        
# Usage
#doc_path = './input_docx/P6700 ITM Vol 1 - Intro, Scope, Competition.docx'
#doc_path = './input_docx/P6700 ITM Vol 2 - Design, Market, and Values.docx'
#doc_path = './input_docx/P6700 ITM Vol 3 - Feature Targets, Testing.docx'
#doc_path = './input_docx/P6700 ITM Vol 4 - Platform, Models, Packages, Adaptation.docx' 
#doc_path = './input_docx/P6700 ITM Vol 5 - Advanced Driver Assistance Safety Systems.docx' 
doc_path = './input_docx/P6700 ITM Vol 6 - Powertrain, Programmable Features.docx' 
# doc_path = './input_docx/P6700 ITM Vol 7 - Axles, Chassis, Trailer Body Connections, Tires & Wheels.docx' 
# doc_path = './input_docx/P6700 ITM Vol 8 - Cab, Lighting, Audio, Equipment.docx'
# doc_path = './input_docx/P6700 ITM Vol 9 - Business Services, Aftermarket, Warranty.docx' 
#os.path.exists(doc_path)
os.path.exists(doc_path)

#folder_name = './output_docx/P6700 ITM Vol 1 - Intro, Scope, Competition'
#folder_name = './output_docx/P6700 ITM Vol 2 - Design, Market, and Values'
#folder_name = './output_docx/P6700 ITM Vol 3 - Feature Targets, Testing'
#folder_name = './output_docx/P6700 ITM Vol 4 - Platform, Models, Packages, Adaptation' 
#folder_name = './output_docx/P6700 ITM Vol 5 - Advanced Driver Assistance Safety Systems'
folder_name = './output_docx/P6700 ITM Vol 6 - Powertrain, Programmable Features' 
# folder_name = './output_docx/P6700 ITM Vol 7 - Axles, Chassis, Trailer Body Connections, Tires & Wheels' 
# folder_name = './output_docx/P6700 ITM Vol 8 - Cab, Lighting, Audio, Equipment'
# folder_name = './output_docx/P6700 ITM Vol 9 - Business Services, Aftermarket, Warranty' 

os.path.exists(folder_name)

extract_number_heading(doc_path)
save_content_under_headings(doc_path, folder_name)

# Do: save_content_under_headings for all docx in a folder
for file_name in os.listdir('./input_docx/'):
    doc_path = f'./input_docx/{file_name}'
    folder_name = f'./output_docx/{file_name.split(".")[0]}'
    save_content_under_headings(doc_path, folder_name)
