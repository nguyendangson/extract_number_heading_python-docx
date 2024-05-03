# import docx

# #open docx file with python-docx
# document = docx.Document("exp/docx_1.docx")

# # #extract body elements
# # body_elements = document._body._body
# # #extract those wrapped in <w:r> tag
# # rs = body_elements.xpath('.//w:r')
# # #check if style is hyperlink (toc)
# # table_of_content = [r.text for r in rs if r.style == "Hyperlink"]

# # print(table_of_content)


# all_data = {}

# all_headings = []
# all_headings_under = []

# def iter_headings(paragraphs):
#     curr_index = 0
#     for i, paragraph in enumerate(paragraphs):
#         if paragraph.style.name.startswith('Heading'):
#             all_headings.append(paragraph.text)
#             all_headings_under.append("")
#             curr_index = i
#         elif curr_index!=0:
#             print(curr_index, all_headings_under)
#             all_headings_under[curr_index] += paragraph.text


#             # yield paragraph
#     return all_headings, all_headings_under

# iter_headings(document.paragraphs)

# print(all_headings)
# # print(all_data.keys())       

# # for heading in iter_headings(document.paragraphs):
# #     # print (heading.text)
# #     continue



from docx import Document
import os

def save_content_under_headings(doc_path, folder_name):
    # Load the Word document
    doc = Document(doc_path)
    
    # Create a new folder
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    
    current_heading = None
    content = []
    current_doc = None
    
    doc_name = doc_path.split("/")[-1]
    doc_name = doc_name.split(".")[0:-1]

    for paragraph in doc.paragraphs:
        # Check if the paragraph is a heading
        if paragraph.style.name.startswith('Heading'):
            # Save the previous section if there is one
            if current_heading and current_doc:
                file_path = os.path.join(folder_name, f"{doc_name}{current_heading}.docx")
                # with open(file_path, 'w', encoding='utf-8') as f:
                #     f.write('\n'.join(content))
                current_doc.save(file_path)
            
            # Reset for the next section
            current_heading = paragraph.text.replace('/', '-').replace('\\', '-')
            content = []
            current_doc = Document()
            current_doc.add_heading(paragraph.text, level=0)
        else:
            # Accumulate content under the current heading
            # content.append(paragraph.text)
            if current_doc is not None:
                current_doc.add_paragraph(paragraph.text)
    
    # Don't forget to save the last section
    if current_heading and current_doc:
        file_path = os.path.join(folder_name, f"{doc_name}{current_heading}.docx")
        current_doc.save(file_path)
        # with open(file_path, 'w', encoding='utf-8') as f:
        #     f.write('\n'.join(content))

# Usage
doc_path = 'exp/docx_1.docx'
folder_name = 'separated_content_docx'
save_content_under_headings(doc_path, folder_name)





